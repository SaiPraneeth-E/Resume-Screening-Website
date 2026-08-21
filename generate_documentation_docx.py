import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, color=RGBColor(15, 23, 42)):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.bold = True
    return p

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_row = table.rows[0]
    hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    for idx, header_text in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "0F2942")  # Deep navy
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_value))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 41, 59)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    # Add spacing after table
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(8)
    return table

def add_callout_box(doc, text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F9FF")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(2, 132, 199)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def generate_doc():
    doc = Document()

    # Set standard page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # -------------------------------------------------------------
    # COVER / TITLE SECTION
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("SMART RESUME SCREENER & CANDIDATE INTELLIGENCE PLATFORM")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(15, 41, 66)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(16)
    sub_run = sub_p.add_run("Comprehensive Technical Design, Architecture, AI Matching Engine & Deployment Report")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    # Metadata Table
    meta_headers = ["Project Attribute", "Details & References"]
    meta_data = [
        ["Author / Lead Developer", "Edupulapati Sai Praneeth"],
        ["Contact Email", "edupulapatisairpaneeth12345@gmail.com"],
        ["Live Production Website", "https://smart-resume-screener-frontend-c1ec.onrender.com"],
        ["Live Backend API Endpoint", "https://smart-resume-screener-backend-69ss.onrender.com"],
        ["GitHub Source Repository", "https://github.com/SaiPraneeth-E/Resume-Screening-Website"],
        ["Tech Stack Overview", "FastAPI (Python 3.10+), React 18, TypeScript, Tailwind CSS, PostgreSQL, Scikit-Learn"],
        ["Documentation Version", "1.0.0 (Production Release)"],
        ["Release Date", "August 2026"]
    ]
    create_styled_table(doc, meta_headers, meta_data, [Inches(2.2), Inches(4.3)])

    add_callout_box(
        doc,
        "This document details the end-to-end technical specification, semantic resume parsing algorithms, candidate scoring workflows, and cloud deployment pipelines developed for the Smart Resume Screener application.",
        title="DOCUMENT PURPOSE"
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # 1. EXECUTIVE SUMMARY & OBJECTIVES
    # -------------------------------------------------------------
    add_styled_heading(doc, "1. Executive Summary & Problem Statement", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "Recruiting and talent acquisition teams receive hundreds of resumes for every posted job description. "
        "Manual resume screening is notoriously labor-intensive, vulnerable to human fatigue and unconscious bias, and frequently "
        "slows down the hiring cycle. Traditional keyword-based Applicant Tracking Systems (ATS) often reject highly qualified candidates "
        "due to rigid string matching (e.g., failing to equate 'React.js' with 'React' or 'PostgreSQL' with 'Postgres')."
    )

    p2 = doc.add_paragraph()
    p2.add_run(
        "Smart Resume Screener is an enterprise-grade AI-powered web platform created by "
    )
    r_author = p2.add_run("Edupulapati Sai Praneeth")
    r_author.bold = True
    p2.add_run(
        " to solve these exact challenges. The platform parses resumes in PDF and TXT formats, normalizes technical skills through an extensible "
        "taxonomy, and computes multidimensional match scores across seven distinct evaluation vectors. It delivers interactive analytics, "
        "customized interview questions, personalized outreach email drafts, and instant side-by-side candidate comparisons."
    )

    # -------------------------------------------------------------
    # 2. KEY FEATURES & CAPABILITIES
    # -------------------------------------------------------------
    add_styled_heading(doc, "2. Key Features & Capabilities", level=1)

    features = [
        ("High-Precision Resume Parsing", "Extracts structured text from complex multi-column PDF and TXT documents using PyMuPDF (fitz), identifying contact details, work experience, education, projects, skills, and certifications."),
        ("Skill Taxonomy & Synonym Normalization", "Maps over 200+ technology and tool aliases to standard taxonomy entries (e.g., 'NodeJS' -> 'Node.js', 'K8s' -> 'Kubernetes', 'ML' -> 'Machine Learning')."),
        ("7-Factor Composite Scoring Engine", "Evaluates candidate fit using weighted metrics: Skill Overlap (35%), Semantic Fit (25%), Work Experience (15%), Project Relevance (10%), Education (5%), Certifications (5%), and Keyword Density (5%)."),
        ("AI-Generated Strengths & Gaps Analysis", "Highlights proven candidate qualifications alongside missing requirements and provides actionable suggestions for resume optimization."),
        ("Automated Interview Kit Generation", "Automatically synthesizes targeted technical, behavioral, and architectural interview questions tailored specifically to the candidate's documented experience and identified skill gaps."),
        ("One-Click Recruiter Outreach Drafter", "Generates professional email communication templates addressed to candidates with personalized praise of their specific projects and background."),
        ("Side-by-Side Candidate Comparison Matrix", "Enables recruiters to select multiple candidates and inspect their metrics, radar charts, and skill breakdowns in a unified comparative view."),
        ("Visual Analytics & Executive Dashboard", "Provides hiring managers with candidate score distribution curves, skill demand charts, shortlist metrics, and screening session history.")
    ]

    for title, desc in features:
        bp = doc.add_paragraph(style='List Bullet')
        r_bt = bp.add_run(f"{title}: ")
        r_bt.bold = True
        r_bt.font.color.rgb = RGBColor(15, 41, 66)
        bp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 3. TECHNICAL ARCHITECTURE & STACK
    # -------------------------------------------------------------
    add_styled_heading(doc, "3. System Architecture & Technology Stack", level=1)

    p_arch = doc.add_paragraph()
    p_arch.add_run(
        "The application is engineered with a decoupled client-server architecture ensuring high responsiveness, strict data validation, "
        "and modular extensibility."
    )

    stack_headers = ["Layer / Domain", "Technologies Used", "Key Architectural Responsibility"]
    stack_data = [
        ["Frontend UI/UX", "React 18, TypeScript, Vite, Tailwind CSS, Lucide Icons, Recharts", "Cyberpunk-inspired dark theme, dynamic radar charts, responsive candidate matrices, and resume drag-and-drop."],
        ["Backend REST API", "Python 3.10+, FastAPI, Uvicorn, Pydantic v2, Asyncio", "High-performance asynchronous request handling, multipart upload parsing, CORS management, and RESTful routing."],
        ["Database Layer", "SQLAlchemy 2.0 (Async), PostgreSQL (Production), SQLite (Dev)", "Relational data persistence for jobs, candidate profiles, screening sessions, scores, and shortlists."],
        ["Document Parsing", "PyMuPDF (fitz), Regex Pattern Matching", "Stream extraction of text chunks, email/phone/link regex resolution, and section segmentation."],
        ["AI & Matcher Engine", "Scikit-Learn (TF-IDF & Cosine Similarity), Sentence-Transformers, NumPy", "Zero-OOM, sub-millisecond semantic document embedding and multi-vector compatibility scoring."],
        ["Deployment & Hosting", "Render (Blueprint YAML), Docker, PostgreSQL Cloud DB, Vercel", "Automated CI/CD git-backed builds, environment injection, and managed database provisioning."]
    ]
    create_styled_table(doc, stack_headers, stack_data, [Inches(1.5), Inches(2.3), Inches(2.7)])

    doc.add_page_break()

    # -------------------------------------------------------------
    # 4. SCREENING ENGINE & MATHEMATICAL FORMULATION
    # -------------------------------------------------------------
    add_styled_heading(doc, "4. Multi-Factor Scoring Engine", level=1)

    p_math = doc.add_paragraph()
    p_math.add_run(
        "To prevent misleading scores caused by single-metric evaluations, the screening engine evaluates candidates across seven independent vectors. "
        "The overall candidate score S_total is calculated as follows:"
    )

    p_formula = doc.add_paragraph()
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_form = p_formula.add_run(
        "S_total = (0.35 * S_skill) + (0.25 * S_semantic) + (0.15 * S_exp) + (0.10 * S_proj) + (0.05 * S_edu) + (0.05 * S_cert) + (0.05 * S_kw)"
    )
    r_form.bold = True
    r_form.font.name = 'Consolas'
    r_form.font.size = Pt(11)
    r_form.font.color.rgb = RGBColor(2, 132, 199)

    score_headers = ["Evaluation Factor", "Weight", "Algorithm & Evaluation Methodology"]
    score_data = [
        ["Skill Match (S_skill)", "35%", "Exact & normalized taxonomy matching. 75% weight on required skills ratio + 25% weight on preferred skills ratio."],
        ["Semantic Fit (S_semantic)", "25%", "TF-IDF / SentenceTransformer Cosine similarity comparing full resume text against complete job responsibilities and brief."],
        ["Experience (S_exp)", "15%", "Calculated from documented role count and job title alignment against target position level."],
        ["Projects (S_proj)", "10%", "Portfolio project volume and intersection of project technologies with required technical competencies."],
        ["Education (S_edu)", "5%", "Degree classification: Master's/Ph.D. (100%), Bachelor's/B.Tech (85%), Other/Diploma (70%)."],
        ["Certifications (S_cert)", "5%", "Presence of verified industry credentials (AWS, Azure, GCP, PMP, Kubernetes, etc.)."],
        ["Keywords (S_kw)", "5%", "Density of target job description keywords found within the body of the candidate's resume."]
    ]
    create_styled_table(doc, score_headers, score_data, [Inches(1.8), Inches(0.9), Inches(3.8)])

    # -------------------------------------------------------------
    # 5. VISUAL WALKTHROUGH & SCREENSHOTS
    # -------------------------------------------------------------
    add_styled_heading(doc, "5. User Interface & Visual Walkthrough", level=1)

    p_screens = doc.add_paragraph()
    p_screens.add_run(
        "The user interface incorporates high-contrast cyber-aesthetic design, smooth glassmorphism effects, vibrant status indicators, "
        "and rich data visualizations built with Tailwind CSS and Recharts."
    )

    screenshots = [
        ("docs/screenshots/landing_page.png", "Figure 1: Smart Resume Screener Landing Page & Feature Showcase", "Interactive landing page introducing recruiters to the AI engine, instant metrics, and workflow features."),
        ("docs/screenshots/screen_page.png", "Figure 2: Resume Screening & Job Specification Input Console", "Job brief editor and multi-file drag-and-drop upload zone for PDF and TXT resume batches."),
        ("docs/screenshots/analytics_page.png", "Figure 3: Analytics & Candidate Intelligence Dashboard", "Comprehensive analytics dashboard displaying score distribution curves, top candidates, and candidate radar metrics.")
    ]

    for img_path, fig_title, fig_desc in screenshots:
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(4)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(0)
            p_cap.paragraph_format.space_after = Pt(12)
            r_ft = p_cap.add_run(fig_title + "\n")
            r_ft.bold = True
            r_ft.font.size = Pt(9.5)
            r_ft.font.color.rgb = RGBColor(15, 41, 66)
            r_fd = p_cap.add_run(fig_desc)
            r_fd.font.italic = True
            r_fd.font.size = Pt(9)
            r_fd.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # -------------------------------------------------------------
    # 6. REST API ENDPOINT SPECIFICATION
    # -------------------------------------------------------------
    add_styled_heading(doc, "6. REST API Endpoint Specification", level=1)

    api_headers = ["HTTP Method", "Endpoint Path", "Functionality & Parameters"]
    api_data = [
        ["GET", "/api/health", "Health check endpoint verifying API status, version, and database connectivity."],
        ["POST", "/api/screen", "Multipart upload endpoint accepting resume files and job specifications. Returns ranked match reports."],
        ["GET", "/api/jobs", "Retrieves all saved job descriptions with title, department, and required skills."],
        ["POST", "/api/jobs", "Creates a new job description with structured skill requirements."],
        ["DELETE", "/api/jobs/{id}", "Deletes a specific job description by UUID."],
        ["GET", "/api/candidates", "Lists screened candidates with query filters: search, min_score, recommendation, shortlisted."],
        ["GET", "/api/candidates/{id}", "Returns detailed candidate profile, parsed experience/education, and match breakdown."],
        ["POST", "/api/candidates/shortlist", "Toggles shortlist bookmark status for a candidate."],
        ["POST", "/api/candidates/compare", "Accepts array of candidate IDs and generates comparative side-by-side matrices."],
        ["GET", "/api/candidates/{id}/interview-questions", "Generates custom interview questions based on candidate skill gaps."],
        ["GET", "/api/candidates/{id}/outreach-email", "Generates personalized recruiter outreach email text."],
        ["GET", "/api/dashboard/stats", "Aggregates total screened resumes, average score, shortlist count, and score tiers."],
        ["GET", "/api/export/csv", "Exports screened candidate rankings and scores to downloadable CSV."]
    ]
    create_styled_table(doc, api_headers, api_data, [Inches(1.2), Inches(2.3), Inches(3.0)])

    # -------------------------------------------------------------
    # 7. LOCAL SETUP & INSTALLATION GUIDE
    # -------------------------------------------------------------
    add_styled_heading(doc, "7. Installation & Local Development Setup", level=1)

    p_req = doc.add_paragraph()
    p_req.add_run("Prerequisites: Python 3.10+, Node.js 18+, npm 9+.")

    add_styled_heading(doc, "Backend Configuration:", level=2)
    doc.add_paragraph(
        "1. Open terminal and navigate to backend directory:\n"
        "   cd backend\n"
        "2. Create and activate virtual environment:\n"
        "   python -m venv venv\n"
        "   venv\\Scripts\\activate   # On Windows (or source venv/bin/activate on Linux/Mac)\n"
        "3. Install Python dependencies:\n"
        "   pip install -r requirements.txt\n"
        "4. Start the FastAPI development server:\n"
        "   uvicorn app.main:app --host 127.0.0.1 --port 8000 --reload"
    )

    add_styled_heading(doc, "Frontend Configuration:", level=2)
    doc.add_paragraph(
        "1. Open another terminal and navigate to frontend directory:\n"
        "   cd frontend\n"
        "2. Install npm dependencies:\n"
        "   npm install\n"
        "3. Launch the Vite development server:\n"
        "   npm run dev\n"
        "4. Access the web application at http://localhost:5173"
    )

    # -------------------------------------------------------------
    # 8. PRODUCTION DEPLOYMENT ARCHITECTURE
    # -------------------------------------------------------------
    add_styled_heading(doc, "8. Production Cloud Deployment Guide", level=1)

    p_dep = doc.add_paragraph()
    p_dep.add_run(
        "The project is production-ready and configured for continuous deployment using Render Blueprint infrastructure-as-code (render.yaml):"
    )

    dep_headers = ["Service Component", "Hosting Environment", "Configuration & Live URL"]
    dep_data = [
        ["Frontend (Static Site)", "Render Static Web Service", "https://smart-resume-screener-frontend-c1ec.onrender.com\nBuild: npm install && npm run build\nPublish: ./dist"],
        ["Backend (FastAPI)", "Render Python Web Service", "https://smart-resume-screener-backend-69ss.onrender.com\nRuntime: Python 3.10+\nCommand: uvicorn app.main:app --host 0.0.0.0 --port $PORT"],
        ["Database", "Render Managed PostgreSQL", "PostgreSQL database with asyncpg connection pooling & auto schema initialization."],
        ["Containerization", "Docker & Docker Compose", "docker-compose up --build -d launches frontend, backend, and PostgreSQL in isolated containers."]
    ]
    create_styled_table(doc, dep_headers, dep_data, [Inches(1.8), Inches(1.8), Inches(2.9)])

    # -------------------------------------------------------------
    # 9. CONCLUSION & AUTHOR INFORMATION
    # -------------------------------------------------------------
    add_styled_heading(doc, "9. Conclusion & Project Sign-Off", level=1)

    p_conc = doc.add_paragraph()
    p_conc.add_run(
        "Smart Resume Screener successfully bridges the gap between raw unstructured candidate resumes and actionable hiring intelligence. "
        "By leveraging semantic text similarity, multi-factor scoring algorithms, automated interview question generators, and an intuitive cyber-themed UI, "
        "the application drastically accelerates recruitment workflows while promoting consistent, fair, and evidence-based candidate evaluations."
    )

    add_callout_box(
        doc,
        "Project Developer: Edupulapati Sai Praneeth\n"
        "GitHub Profile: https://github.com/SaiPraneeth-E\n"
        "Repository: https://github.com/SaiPraneeth-E/Resume-Screening-Website\n"
        "Email Contact: edupulapatisairpaneeth12345@gmail.com\n"
        "Live Application: https://smart-resume-screener-frontend-c1ec.onrender.com",
        title="DEVELOPER CREDENTIALS & LINKS"
    )

    output_path = os.path.abspath("Smart_Resume_Screener_Project_Report.docx")
    doc.save(output_path)
    print(f"Document successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_doc()
